import os
import uuid
from typing import List, Dict, Any
from datetime import datetime
import asyncio
from docx import Document
from docx.shared import Inches
import json

from config import CUSTOMER_ID, FILES_DIR
from db.qdrant_service import get_qdrant_client
from services.embeddings_service import ensure_cosine_collection, get_embeddings_from_llama, save_embeddings_with_path
from services.file_service import normalize_target_path
from services.local_llma_service import LocalLlamaService
from services.projects_handler import get_project_file, get_project_file_content
from services.templates_service import TEMPLATES_BASE_DIR, get_template_metadata
from services.utils import extract_text_from_file

# async def generate_filled_template(
#     template_file_id: str,
#     source_file_ids: List[str],
#     project_name: str = None
# ) -> Dict[str, Any]:
#     """
#     Generate a filled template by combining template structure with source file content
    
#     Args:
#         template_file_id: ID of the template file in Qdrant (from templates collection)
#         source_file_ids: List of source file IDs in Qdrant (from documents collection)
#         project_name: Project name for organizing output files
    
#     Returns:
#         Dictionary with generation results and file info
#     """
#     try:
#         print(f"🔄 Starting template generation for template: {template_file_id}")
#         print(f"   Source files: {source_file_ids}")
#         print(f"   Project: {project_name}")
        
#         # Get template metadata and content (from templates collection)
#         template_metadata = await get_template_metadata(template_file_id)
#         if not template_metadata:
#             raise ValueError(f"Template file not found: {template_file_id}")
        
#         template_path = template_metadata.get("file_path")
#         if not os.path.exists(template_path):
#             raise ValueError(f"Template file not found on disk: {template_path}")
        
#         # Get source files metadata and content (from documents collection)
#         source_files_content = []
#         for source_file_id in source_file_ids:
#             # Try to get source file from documents collection first
#             source_metadata = await get_project_file(source_file_id)
            
#             # If not found in documents collection, try templates collection as fallback
#             if not source_metadata:
#                 print(f"⚠️  Source file not found in documents collection: {source_file_id}, trying templates collection...")
#                 source_metadata = await get_template_metadata(source_file_id)
            
#             if not source_metadata:
#                 print(f"⚠️  Source file not found in any collection: {source_file_id}, skipping")
#                 continue
            
#             # Try to get file path from different possible field names
#             source_path = (
#                 source_metadata.get("file_path") or 
#                 source_metadata.get("full_file_path") or
#                 source_metadata.get("relative_path")
#             )
            
#             if source_path and os.path.exists(source_path):
#                 content = extract_text_from_file(source_path)
#                 source_files_content.append({
#                     "filename": source_metadata.get("filename") or source_metadata.get("original_filename", "Unknown"),
#                     "content": content,
#                     "metadata": source_metadata
#                 })
#                 print(f"📄 Loaded source: {source_metadata.get('filename') or source_metadata.get('original_filename', 'Unknown')} ({len(content)} chars)")
#             else:
#                 print(f"⚠️  Source file not on disk: {source_path}")
#                 # Try to get content from Qdrant chunks if file not on disk
#                 content = await get_project_file_content(source_file_id)
#                 if content:
#                     source_files_content.append({
#                         "filename": source_metadata.get("filename") or source_metadata.get("original_filename", "Unknown"),
#                         "content": content,
#                         "metadata": source_metadata
#                     })
#                     print(f"📄 Loaded source from Qdrant: {source_metadata.get('filename') or source_metadata.get('original_filename', 'Unknown')} ({len(content)} chars)")
        
#         if not source_files_content:
#             raise ValueError("No valid source files found to process")
        
#         # Read template content
#         template_content = extract_text_from_file(template_path)
#         print(f"📋 Template content loaded: {len(template_content)} chars")
        
#         # Use LLM to fill the template
#         filled_content = await fill_template_with_llm(
#             template_content=template_content,
#             source_files_content=source_files_content,
#             template_metadata=template_metadata
#         )
        
#         # Create the filled document
#         output_file_info = await create_filled_document(
#             filled_content=filled_content,
#             template_metadata=template_metadata,
#             source_files_metadata=[s["metadata"] for s in source_files_content],
#             project_name=project_name
#         )
        
#         print(f"✅ Template generation completed successfully")
#         print(f"   Output file: {output_file_info['file_path']}")
        
#         return {
#             "success": True,
#             "template_file_id": template_file_id,
#             "source_file_ids": source_file_ids,
#             "output_file": output_file_info,
#             "generated_at": datetime.now().isoformat(),
#             "project_name": project_name
#         }
        
#     except Exception as e:
#         print(f"❌ Template generation failed: {e}")
#         import traceback
#         print(f"🔍 Stack trace: {traceback.format_exc()}")
#         return {
#             "success": False,
#             "error": str(e),
#             "template_file_id": template_file_id,
#             "source_file_ids": source_file_ids,
#             "generated_at": datetime.now().isoformat()
#         }
    
    
async def fill_template_with_llm(
    template_file_path: str,
    source_files_content: List[Dict[str, Any]],
    template_metadata: Dict[str, Any]
) -> str:
    """
    Use LLM to fill the actual template file while preserving structure
    """
    # Read the actual template file content
    print(f"🔍 DEBUG: Reading template file from: {template_file_path}")
    
    template_content = extract_text_from_file(template_file_path)
    
    # Prepare context from source files with debugging
  
    
    context_parts = []
    for i, source in enumerate(source_files_content):

        context_parts.append(f"=== SOURCE FILE {i+1}: {source['filename']} ===")
        context_parts.append(source['content'])
        context_parts.append("")  # Empty line between sources
    
    context = "\n".join(context_parts)
    
    # Get template filename for reference
    template_name = template_metadata.get("original_filename", "template")
    
    # Optimized prompt using actual template structure
    prompt = f"""
ACTUAL TEMPLATE FILE STRUCTURE (fill this exact structure):
{template_content}
SOURCE DATA (use to fill the template):
{context}

CRITICAL INSTRUCTIONS:
1. PRESERVE THE EXACT STRUCTURE from the template above
2. Keep all original headings, sections, and labels
3. Fill tables with appropriate data from sources
4. Use the same table format as shown in the template
5. Only replace placeholder text and fill empty fields
6. Maintain all original formatting and styling cues
7. If information is missing, use "TBD" or leave as-is
8. Do not add any new sections or content
9. if you see [TABLE:] means the template contains at this position a table. remove that word [TABLE:] and fill the table 

OUTPUT REQUIREMENTS:
- Return the complete filled template
- Maintain identical structure to the original
- Use the same table formats and headings
- No additional explanations or content
"""
    
    print(f"🤖 whole prompt is:")
    print(f"📋{prompt}")
   
  
    # Use the local LLM service
    print(f"🚀 Calling LLM...")
    
    # Use the local LLM service
    filled_content = await local_llama.get_chat_completion_with_fullprompt(prompt)
    
    print(f"✅ LLM generation completed: {len(filled_content)} chars")
        
    # Print first 500 chars of LLM output for debugging
    print(f"🔍 DEBUG: First 500 chars of LLM output:")
    print(filled_content[:500])
    print("...")
    
    return filled_content

async def create_filled_document(
    filled_content: str,
    template_metadata: Dict[str, Any],
    source_files_metadata: List[Dict[str, Any]],
    project_name: str
) -> Dict[str, Any]:
    """
    Create a Word document with the exact content generated by LLM
    """
    file_path = None  # Initialize file_path here
    
    try:
        # Validate project name is provided
        if not project_name or project_name.strip() == "":
            raise ValueError("Project name must be specified")
        
        # Normalize target path
        target_path = "generated"
        normalized_target_path = normalize_target_path(target_path)
        
        # Create full target directory
        full_target_dir = os.path.join(FILES_DIR, project_name, normalized_target_path)
        os.makedirs(full_target_dir, exist_ok=True)
        
        # Generate output filename
        template_name = template_metadata.get("original_filename", "template")
        base_name = os.path.splitext(template_name)[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{base_name}_filled_{timestamp}.docx"
        
        # Generate file ID and create file path
        file_id = str(uuid.uuid4())
        file_path = os.path.join(full_target_dir, f"{file_id}_{output_filename}")
        relative_file_path = os.path.join(normalized_target_path, f"{file_id}_{output_filename}")
        
        # Create document with LLM content exactly as generated
        create_document_from_llm_content(file_path, filled_content)
        
        # Get file info
        file_size = os.path.getsize(file_path)
        created_time = datetime.now()
        
        # Process the file for embeddings
        text_for_embedding = filled_content[:5000]
        if not text_for_embedding or len(text_for_embedding.strip()) < 10:
            raise ValueError("No readable text content found in generated document")
        
        chunks = [text_for_embedding[i:i + 500] for i in range(0, len(text_for_embedding), 500)]
        chunks = [chunk for chunk in chunks if len(chunk.strip()) > 10]
        
        if not chunks:
            raise ValueError("No valid text chunks created from generated document")
        
        print(f"📊 Processing generated document: {len(chunks)} chunks")
        
        # Get embeddings
        embeddings = get_embeddings_from_llama(chunks)
        if not embeddings or len(embeddings) != len(chunks):
            raise ValueError("Embedding count mismatch for generated document")
        
        # Save to Qdrant
        qdrant_client = get_qdrant_client()
        collection_name = f"customer_{CUSTOMER_ID}_documents"
        ensure_cosine_collection(qdrant_client, collection_name, vector_size=4096)
        
        save_embeddings_with_path(
            collection_name=collection_name,
            file_id=file_id,
            filename=output_filename,
            chunks=chunks,
            embeddings=embeddings,
            target_path=normalized_target_path,
            target_project=project_name,
            auto_generated=True,
            source_template_id=template_metadata.get("file_id")
        )
        
        print(f"💾 Saved generated document: {file_path}")
        print(f"🗄️  Metadata saved with ID: {file_id}")
        
        return {
            "fileId": file_id,
            "filename": output_filename,
            "savedAs": f"{file_id}_{output_filename}",
            "path": normalized_target_path,
            "project": project_name,
            "fullPath": relative_file_path,
            "size": file_size,
            "fileType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "created": created_time.isoformat(),
            "chunks": len(chunks),
            "metadataSaved": True,
            "autoGenerated": True,
            "sourceTemplateId": template_metadata.get("file_id")
        }
        
    except Exception as e:
        print(f"❌ Error creating filled document: {e}")
        import traceback
        print(f"🔍 Stack trace: {traceback.format_exc()}")
        
        # Clean up file if it was created but processing failed
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                print(f"🧹 Cleaned up failed file: {file_path}")
            except Exception as cleanup_error:
                print(f"⚠️  Could not clean up file {file_path}: {cleanup_error}")
        
        raise


def create_document_from_llm_content(file_path: str, filled_content: str):
    """
    Create Word document with LLM content exactly as generated
    """
    doc = Document()
    
    # Split content by lines and process each line
    lines = filled_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()  # Add empty paragraph for blank lines
            continue
            
        # Check if this line looks like table data (pipe-separated or tabular)
        if '|' in line and line.count('|') >= 2:
            # This is a table row - parse and add as table
            add_table_row(doc, line)
        elif line.startswith('---') or line.startswith('==='):
            # This is a table separator - skip or handle as needed
            continue
        elif looks_like_table_header(line):
            # This might be a table header
            add_table_header(doc, line)
        else:
            # Regular text - add as paragraph
            doc.add_paragraph(line)
    
    doc.save(file_path)
    print(f"✅ Created document with LLM content: {file_path}")


def looks_like_table_header(line: str) -> bool:
    """Check if a line looks like a table header"""
    return any(indicator in line.lower() for indicator in [
        'ref', 'action', 'who', 'when', 'item', 'decision', 
        'responsible', 'deadline', 'description'
    ])


def add_table_row(doc: Document, line: str):
    """Add a table row from pipe-separated data"""
    # Remove leading/trailing pipes and split
    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
    
    if len(cells) >= 2:
        # Check if we need to create a new table or add to existing one
        if not hasattr(doc, '_current_table') or doc._current_table is None:
            # Create new table
            doc._current_table = doc.add_table(rows=1, cols=len(cells))
            doc._current_table.style = 'Table Grid'
        else:
            # Add row to existing table
            doc._current_table.add_row()
        
        # Add cells to the current row
        current_row = doc._current_table.rows[-1]
        for i, cell_content in enumerate(cells):
            if i < len(current_row.cells):
                current_row.cells[i].text = cell_content


def add_table_header(doc: Document, line: str):
    """Add a table header"""
    # For now, just add as a heading - the LLM should generate proper table format
    doc.add_heading(line, level=2)

async def fill_template_document(
    template_metadata: Dict[str, Any],
    structured_data: Dict[str, Any],
    output_path: str
):
    """
    Fill a template document with structured data while preserving original structure
    """
    template_path = template_metadata.get("file_path")
    
    if not os.path.exists(template_path):
        raise ValueError(f"Template file not found: {template_path}")
    
    # Copy the template to output path
    import shutil
    shutil.copy2(template_path, output_path)
    
    # Open the copied document
    doc = Document(output_path)
    
    # Replace placeholders throughout the document
    replace_placeholders_in_document(doc, structured_data)
    
    # Save the modified document
    doc.save(output_path)


def replace_placeholders_in_document(doc: Document, structured_data: Dict[str, Any]):
    """
    Replace placeholders in the document with actual data while preserving structure
    """
    # Common placeholder mappings
    placeholder_map = {
        # Meeting details
        '[Project Name]': structured_data.get('project_name', ''),
        '[PROJECT_NAME]': structured_data.get('project_name', ''),
        '[Date]': structured_data.get('date', ''),
        '[DATE]': structured_data.get('date', ''),
        '[Time]': structured_data.get('time', ''),
        '[TIME]': structured_data.get('time', ''),
        '[Location]': structured_data.get('location', ''),
        '[LOCATION]': structured_data.get('location', ''),
        '[Purpose]': structured_data.get('purpose', ''),
        '[PURPOSE]': structured_data.get('purpose', ''),
        
        # Project info
        '[Budget]': structured_data.get('budget', ''),
        '[BUDGET]': structured_data.get('budget', ''),
        '[Timeline]': structured_data.get('timeline', ''),
        '[TIMELINE]': structured_data.get('timeline', ''),
        
        # Generic placeholders
        '[TBD]': 'To be determined',
        '[TODO]': 'To be determined',
        '...': 'Not specified',
    }
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        new_text = original_text
        
        for placeholder, value in placeholder_map.items():
            if placeholder in original_text and value:
                new_text = new_text.replace(placeholder, value)
        
        if new_text != original_text:
            # Clear and rebuild the paragraph to preserve runs and formatting
            paragraph.clear()
            # Add the new text while trying to preserve basic formatting
            if any(keyword in new_text.lower() for keyword in ['project', 'date', 'time', 'location']):
                # Likely a heading or important field
                run = paragraph.add_run(new_text)
                run.bold = True
            else:
                paragraph.add_run(new_text)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                new_text = original_text
                
                for placeholder, value in placeholder_map.items():
                    if placeholder in original_text and value:
                        new_text = new_text.replace(placeholder, value)
                
                if new_text != original_text:
                    cell.text = new_text
    
    # Fill table data if we have structured content
    fill_table_data(doc, structured_data)


def fill_table_data(doc: Document, structured_data: Dict[str, Any]):
    """
    Fill table data based on content structure
    """
    for table in doc.tables:
        table_text = ' '.join(cell.text for row in table.rows for cell in row.cells).lower()
        
        # Identify table type and fill accordingly
        if any(keyword in table_text for keyword in ['action', 'who', 'when']):
            fill_generic_table(table, structured_data.get('actions', []), 'actions')
        elif any(keyword in table_text for keyword in ['item', 'discuss']):
            fill_generic_table(table, structured_data.get('items', []), 'items')
        elif any(keyword in table_text for keyword in ['decision']):
            fill_generic_table(table, structured_data.get('decisions', []), 'decisions')
        elif any(keyword in table_text for keyword in ['attend', 'participant']):
            fill_generic_table(table, structured_data.get('attendees', []), 'attendees')


def fill_generic_table(table, data_items: List[str], data_type: str):
    """
    Fill a generic table with data items, preserving the original structure
    """
    if not data_items:
        return
    
    # Determine how many data rows we need (excluding header)
    available_rows = len(table.rows) - 1  # Assume first row is header
    
    if available_rows <= 0:
        return
    
    # Fill available rows with data
    for i in range(min(available_rows, len(data_items))):
        row_idx = i + 1  # Start after header
        if row_idx < len(table.rows):
            # Fill first cell with reference number
            if len(table.rows[row_idx].cells) > 0:
                table.rows[row_idx].cells[0].text = str(i + 1)
            
            # Fill second cell with the data item
            if len(table.rows[row_idx].cells) > 1:
                table.rows[row_idx].cells[1].text = data_items[i]
            
            # Fill additional cells if they exist and we have more data
            if data_type == 'actions' and len(table.rows[row_idx].cells) > 2:
                # For actions, fill "Who" and "By When" with placeholders
                if len(table.rows[row_idx].cells) > 2:
                    table.rows[row_idx].cells[2].text = 'TBD'
                if len(table.rows[row_idx].cells) > 3:
                    table.rows[row_idx].cells[3].text = 'TBD'


def parse_llm_content_to_structured_data(filled_content: str) -> Dict[str, Any]:
    """
    Parse LLM content into structured data for template filling
    """
    structured_data = {
        "project_name": "",
        "date": "",
        "time": "",
        "location": "",
        "purpose": "",
        "budget": "",
        "timeline": "",
        "items": [],
        "actions": [],
        "decisions": [],
        "attendees": []
    }
    
    lines = filled_content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Extract basic project info
        if line.startswith('**Project:**'):
            structured_data["project_name"] = line.replace('**Project:**', '').strip()
        elif line.startswith('**Date:**'):
            structured_data["date"] = line.replace('**Date:**', '').strip()
        elif 'Budget:' in line:
            structured_data["budget"] = extract_value_after_colon(line)
        elif 'Timeline:' in line:
            structured_data["timeline"] = extract_value_after_colon(line)
        
        # Categorize content based on context
        if line.startswith('* '):
            item = line.replace('*', '').strip()
            if not current_section:
                # Auto-categorize based on keywords
                if any(keyword in item.lower() for keyword in ['develop', 'implement', 'create', 'build', 'complete']):
                    structured_data["actions"].append(item)
                elif any(keyword in item.lower() for keyword in ['agree', 'decide', 'approve', 'conclude']):
                    structured_data["decisions"].append(item)
                else:
                    structured_data["items"].append(item)
            else:
                structured_data[current_section].append(item)
    
    return structured_data


def extract_value_after_colon(line: str) -> str:
    """Extract value after colon, handling markdown formatting"""
    if ':' in line:
        return line.split(':', 1)[1].replace('**', '').strip()
    return ""


    
    # Global instance
local_llama = LocalLlamaService()




