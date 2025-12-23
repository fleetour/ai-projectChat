import asyncio
from docx import Document
import numpy as np
from typing import Any, Dict, List, Literal, Optional
import subprocess
import os
import logging
import tempfile

logger = logging.getLogger(__name__)

from config import CUSTOMER_ID

def get_content_type(filename: str) -> str:
    """Determine content type based on file extension"""
    import mimetypes
    mimetypes.init()
    
    # Common extensions mapping
    extension_map = {
        '.pdf': 'application/pdf',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xls': 'application/vnd.ms-excel',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.ppt': 'application/vnd.ms-powerpoint',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        '.txt': 'text/plain',
        '.csv': 'text/csv',
        '.json': 'application/json',
        '.xml': 'application/xml',
        '.zip': 'application/zip',
        '.rar': 'application/x-rar-compressed',
        '.7z': 'application/x-7z-compressed',
        '.tar': 'application/x-tar',
        '.gz': 'application/gzip',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.svg': 'image/svg+xml',
        '.mp3': 'audio/mpeg',
        '.mp4': 'video/mp4',
        '.avi': 'video/x-msvideo',
        '.mov': 'video/quicktime',
        '.py': 'text/x-python',
        '.js': 'application/javascript',
        '.html': 'text/html',
        '.css': 'text/css'
    }
    
    # Get file extension
    _, ext = os.path.splitext(filename.lower())
    
    # Check our mapping first
    if ext in extension_map:
        return extension_map[ext]
    
    # Fallback to mimetypes
    content_type, _ = mimetypes.guess_type(filename)
    return content_type or "application/octet-stream"

def get_collection_name(collection_type: Literal["documents", "templates", "conversations"], customer_id: str) -> str:
    """
    Get the full collection name based on collection type
    
    Args:
        collection_type: Either "documents", "templates", or "conversations"
        
    Returns:
        Full collection name with customer prefix
    """
    if collection_type not in ["documents", "templates", "conversations"]:
        raise ValueError(f"Invalid collection type: {collection_type}. Must be 'documents', 'templates', or 'conversations'")
    
    return f"customer_{customer_id}_{collection_type}"

def normalize_vector(vector: List[float]) -> List[float]:
    """Normalize vector for cosine similarity"""
    vector_array = np.array(vector)
    norm = np.linalg.norm(vector_array)
    if norm == 0:
        return vector  # Return as-is if zero vector
    normalized = vector_array / norm
    return normalized.tolist()

def validate_file_type(filename: str) -> None:
    """
    Validate if file type is supported. Raises ValueError for unsupported types.
    """
    SUPPORTED_EXTENSIONS = {'.docx', '.pdf', '.txt'}
    
    file_extension = os.path.splitext(filename.lower())[1]
    
    if not file_extension:
        raise ValueError(f"File '{filename}' has no extension. Supported types: {', '.join(SUPPORTED_EXTENSIONS)}")
    
    if file_extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{file_extension}' for file '{filename}'. "
            f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}. "
            f"Please convert to a supported format."
        )

async def extract_text_from_bytes_async(content: bytes, filename: str) -> str:
    """Async text extraction."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(
        None,
        extract_text_from_bytes,  # Your existing sync function
        content,
        filename
    )


def extract_text_from_bytes(content: bytes, filename: str, original_file_path: Optional[str] = None) -> str:
    """
    Extract text from file bytes by saving to temp file and using existing extract_text_from_file.
    
    Args:
        content: File content as bytes
        filename: Original filename (for extension detection)
        original_file_path: Optional original path for logging
    
    Returns:
        str: Extracted text
    """
    temp_file_path = None
    
    try:
        # Get file extension from filename
        _, file_ext = os.path.splitext(filename.lower())
        
        # Create a temporary file with the correct extension
        with tempfile.NamedTemporaryFile(
            suffix=file_ext, 
            delete=False,
            prefix="extract_"
        ) as tmp:
            # Write content to temp file
            tmp.write(content)
            temp_file_path = tmp.name
        
        # Use the existing extract_text_from_file function
        text = extract_text_from_file(temp_file_path)
        
        return text
        
    except Exception as e:
        error_context = f" from {original_file_path}" if original_file_path else ""
        logger.error(f"Error extracting text from {filename}{error_context}: {e}")
        
        # Try fallback for common text files
        if file_ext in ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm']:
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return content.decode('latin-1')
                except:
                    return ""
        return ""
    
    finally:
        # Always clean up temp file
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp file {temp_file_path}: {e}")

def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file types including tables in Word documents"""
    ext = file_path.lower()
    
    if ext.endswith(".pdf"):
        from pdfminer.high_level import extract_text
        return extract_text(file_path)
    elif ext.endswith(".docx"):
        from docx import Document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract text from tables
        for table in doc.tables:
            table_text = extract_table_text(table)
            if table_text:
                text_parts.append(table_text)
        
        return "\n".join(text_parts)
    elif ext.endswith(".txt"):
        with open(file_path, "r", encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {file_path}")

def extract_table_text(table) -> str:
    """Extract text from a Word table and format it for LLM understanding"""
    table_lines = []
    
    for row_idx, row in enumerate(table.rows):
        row_cells = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                row_cells.append(cell_text)
        
        if row_cells:
            # Format as pipe-separated table for LLM
            table_lines.append(" | ".join(row_cells))
    
    if table_lines:
        return "[TABLE:]\n" + "\n".join(table_lines)
    return ""


# Optional: Helper function for tag-based searching
def filter_templates_by_tags(templates: List[Dict[str, Any]], required_tags: List[str] = None, 
                           any_tags: List[str] = None, exclude_tags: List[str] = None) -> List[Dict[str, Any]]:
    """
    Filter templates based on tag criteria
    """
    if not any([required_tags, any_tags, exclude_tags]):
        return templates
    
    filtered_templates = []
    
    for template in templates:
        template_tags = template.get("tags", [])
        
        # Check required tags (ALL must be present)
        if required_tags and not all(tag in template_tags for tag in required_tags):
            continue
        
        # Check any tags (AT LEAST ONE must be present)
        if any_tags and not any(tag in template_tags for tag in any_tags):
            continue
        
        # Check exclude tags (NONE can be present)
        if exclude_tags and any(tag in template_tags for tag in exclude_tags):
            continue
        
        filtered_templates.append(template)
    
    return filtered_templates
    
def calculate_adaptive_top_k(query: str) -> int:
    """Determine optimal top_k based on query characteristics."""
    query_lower = query.lower()
    
    # List/Enumerate queries need more context
    list_keywords = ['list', 'all', 'every', 'each', 'attendees', 'members', 
                    'participants', 'who are', 'names of', 'show me']
    
    # Detail queries need more context
    detail_keywords = ['describe', 'explain', 'details', 'summary', 'overview',
                      'what is', 'how does', 'tell me about']
    
    # Simple queries need less context
    simple_keywords = ['yes', 'no', 'when', 'where', 'confirm', 'is there']
    
    for keyword in list_keywords:
        if keyword in query_lower:
            return 10  # More chunks for list queries
            
    for keyword in detail_keywords:
        if keyword in query_lower:
            return 8   # Moderate chunks for detail queries
    
    return 5  # Default for simple queries

def format_results_for_history(results) -> List[Dict]:
    """Format search results for history storage."""
    return [
        {
            "file_id": r.payload.get("file_id"),
            "filename": r.payload.get("filename"),
            "score": float(r.score)
        }
        for r in results[:3]  # Store only top 3 in history
    ]